import pathlib
from pypdf import PdfReader
from docx import Document
import pdfplumber
def extract_text_from_pdf(file_path:str)->str:
    pdf_text=[]
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            pdf_text.append(page.extract_text())
    pdf_text="\n".join(pdf_text)
    return pdf_text

def extract_text_from_docx(file_path:str)->str:
    #doc=Document(file_path) -> Document(obj)
    #doc.paragraphs -> List of Paragraph(objs)
    #doc.paragraphs[0].text -> Content of first paragraph(str)
    doc=Document(file_path)
    doc_text=[]
    for paragraph in doc.paragraphs:
        doc_text.append(paragraph.text)
    doc_text="\n".join(doc_text)
    return doc_text

def extract_text(file_path:str)->str:
    """
    Detects the file type and routes to the correct extractor.
    """
    text=""
    ext=pathlib.Path(file_path).suffix.lower()
    if ext == ".pdf":
        text=extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text=extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    return text


if __name__ == "__main__":
    print(extract_text("testfiles/testpdf.pdf"))
    print("--------------------------------------")
    print(extract_text("testfiles/Resume.pdf"))